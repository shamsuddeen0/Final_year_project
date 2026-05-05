#!/usr/bin/env python3
"""Convert Chapter 1 and 2 markdown files to Microsoft Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

# Create document
doc = Document()

# Set normal style
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Title
title = doc.add_heading('', level=0)
title_run = title.add_run('CHAPTER ONE: INTRODUCTION')
title_run.font.size = Pt(14)
title_run.bold = True

# Read Chapter 1 content
with open('Chapter1_Introduction.md', 'r') as f:
    ch1_content = f.read()

# Read Chapter 2 content
with open('Chapter2_LiteratureReview.md', 'r') as f:
    ch2_content = f.read()

def parse_markdown(content, doc):
    """Parse markdown content and add to document."""
    lines = content.split('\n')
    in_reference = False

    for line in lines:
        line = line.strip()

        # Skip the main title (already added)
        if line.startswith('# CHAPTER'):
            continue

        # Section headings (##)
        if line.startswith('## '):
            heading_text = line.replace('## ', '')
            doc.add_heading(heading_text, level=1)
        # Subsection headings (###)
        elif line.startswith('### '):
            heading_text = line.replace('### ', '')
            doc.add_heading(heading_text, level=2)
        # Handle bold text
        elif '**' in line:
            # Process paragraph with bold formatting
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.+?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.replace('**', ''))
                    run.bold = True
                else:
                    p.add_run(part)
        # Empty line
        elif not line:
            in_reference = False
        # Reference section
        elif line.startswith('## References') or line.startswith('---'):
            in_reference = True
            doc.add_heading('References', level=1)
        # Reference entries
        elif in_reference and line and (line[0].isalpha() or line[0] == '('):
            p = doc.add_paragraph(line, style='List Bullet')
        # Table handling
        elif line.startswith('|'):
            continue  # Skip markdown tables for now
        # Regular paragraph
        elif line and not line.startswith('#'):
            p = doc.add_paragraph(line)
            # Set paragraph justification
            p.alignment = 2  # Justified

# Process Chapter 1
parse_markdown(ch1_content, doc)

# Add page break
doc.add_page_break()

# Add Chapter 2 title
title2 = doc.add_heading('', level=0)
title2_run = title2.add_run('CHAPTER TWO: LITERATURE REVIEW')
title2_run.font.size = Pt(14)
title2_run.bold = True

# Process Chapter 2
parse_markdown(ch2_content, doc)

# Save document
doc.save('Honeyword_IDS_Chapters_1_and_2.docx')
print('Document created: Honeyword_IDS_Chapters_1_and_2.docx')